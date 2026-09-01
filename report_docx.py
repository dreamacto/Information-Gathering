#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Render the template-shaped, report-safe xcx attack result DOCX."""
from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from report_model import aggregate_report_findings, optional_scope_rows
from project_paths import config_path

CST = timezone(timedelta(hours=8))
ROOT = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = ROOT / "templates" / "攻防成果报告_模板.docx"


def now_date() -> str:
    return datetime.now(CST).strftime("%Y-%m-%d")


def load_report_config(path: Path | None) -> dict:
    if path is None:
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        return {}
    value = data.get("reporting", data)
    return value if isinstance(value, dict) else {}


def load_report_policy(path: Path | None) -> dict:
    data = load_report_config(path)
    return data.get("policy", {}) if isinstance(data.get("policy"), dict) else {}


def _setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for i, sz in ((1, 16), (2, 13)):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Arial"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def _clear_body(doc: Document) -> None:
    body = doc._body._element
    for element in list(body):
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def para(doc: Document, value: object = "", *, bold: bool = False, mono: bool = False, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(_sanitize_text(value))
    run.bold = bold
    if mono:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def _values(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, dict):
        return [f"{key}：{value[key]}" for key in value if str(value[key]).strip()]
    return [str(value).strip()] if value and str(value).strip() else []


def _unique(values: list[object]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        text = str(value or "").strip()
        if text and text not in seen:
            seen.add(text)
            result.append(text)
    return result


def _limit(values: list[object], maximum: int = 2) -> list[str]:
    return _unique(values)[:maximum]


def kv_table(doc: Document, rows: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = _sanitize_text("；".join(_values(value)) if isinstance(value, (list, dict)) else str(value or ""))
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True


def _sanitize_text(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"Ins_[A-Za-z0-9]+", "<授权测试账号>", text)
    text = re.sub(r"(?i)(!!vip@instrument!![A-Za-z0-9]+|B7B36F43A9E5806679DD1F0AFD9539D7|INSTRUMENT@123IM\$418)", "<已脱敏密钥>", text)
    text = re.sub(r'(?i)(token|secret|key|sn|session_key|openid|unionid|mobile|phone)([=: ]+)[^&\\s,}]+', r'\1\2<已脱敏>', text)
    text = re.sub(r"(?<![A-Za-z])(?:7741701|Ins4000074077|9261962|967846)(?![A-Za-z])", "<测试对象ID>", text)
    text = re.sub(r"(\+?86[- ]?1\d{10}|1\d{10})", "<测试手机号>", text)
    return text


def _command_text(item: object) -> str:
    if isinstance(item, dict):
        value = str(item.get("cmd") or item.get("command") or item.get("request") or "").strip()
    else:
        value = str(item or "").strip()
    return _sanitize_text(value)



def _render_meta_proofs(doc: Document, meta: dict) -> None:
    para(doc, "资产归属证明网址", bold=True)
    asset_url = meta.get("asset_proof_url")
    para(doc, "；".join(_values(asset_url)) if asset_url else "【请补充资产归属证明网址】")
    filing = meta.get("filing_proof_url")
    if filing or meta.get("include_filing_proof", True):
        para(doc, "备案系统证明网址", bold=True)
        para(doc, "；".join(_values(filing)) if filing else "【请补充备案系统证明网址】")


def _render_reproduction(doc: Document, finding: dict, env_lines: list[str]) -> None:
    para(doc, "详细复现命令或操作步骤", bold=True)
    if env_lines:
        para(doc, "环境准备", bold=True)
        for line in env_lines:
            para(doc, line, mono=True)
    steps = finding.get("steps") or []
    if steps:
        para(doc, "验证步骤", bold=True)
        for index, step in enumerate(steps, 1):
            para(doc, f"{index}. {step}")
    commands = finding.get("reproduction_commands") or finding.get("commands") or []
    if commands:
        para(doc, "真实复现命令", bold=True)
        for command in commands:
            text = _command_text(command)
            if text:
                para(doc, text, mono=True)
    elif not steps:
        para(doc, "真实复现命令", bold=True)
        para(doc, "【请补充实际复现命令】")
    notes = finding.get("notes") or []
    if notes:
        para(doc, "命令备注", bold=True)
        for note in notes:
            para(doc, note)


def _render_result(doc: Document, finding: dict) -> None:
    para(doc, "返回结果/结果解读", bold=True)
    actual = finding.get("actual_result") or ""
    if actual:
        para(doc, "实际结果：" + actual)
    else:
        para(doc, "实际结果：")
    expected = finding.get("expected_result") or ""
    if expected:
        para(doc, "预期/判定依据：" + expected)
    interpretation = finding.get("interpretation") or finding.get("impact") or ""
    if interpretation:
        para(doc, "结果解读：" + interpretation)
    pagination = finding.get("pagination") or []
    if pagination:
        para(doc, "翻页验证", bold=True)
        for item in pagination:
            para(doc, item)
    cleanup = finding.get("cleanup") or []
    if cleanup:
        para(doc, "清理或还原步骤", bold=True)
        for item in cleanup:
            para(doc, item)


def build_report(meta: dict, findings: list[dict], out: Path, report_policy: dict | None = None, template_path: Path | None = None) -> Path:
    report_policy = report_policy or {}
    max_items = int(report_policy.get("max_problems_or_remediations", 2) or 2)
    template = template_path or DEFAULT_TEMPLATE
    doc = Document(template) if template.exists() else Document()
    _clear_body(doc)
    _setup_styles(doc)
    grouped = aggregate_report_findings(findings)

    # The title is fixed by the supplied template; no target/team/date metadata is injected here.
    para(doc, "攻防成果报告", bold=True, style="Title")
    _render_meta_proofs(doc, meta)

    para(doc, "一、目标信息", bold=True, style="Heading 1")
    target_info = meta.get("target_info")
    if target_info:
        rows = [(str(k), v) for k, v in target_info.items()] if isinstance(target_info, dict) else [("目标信息", target_info)]
    elif grouped:
        rows = [("目标系统", grouped[0]["system"]), ("目标URL", "；".join(grouped[0]["urls"]))]
    else:
        rows = [("目标系统", "【请补充目标系统】")]
    kv_table(doc, rows)

    para(doc, "二、成果说明", bold=True, style="Heading 1")
    env_lines = _values(meta.get("env_lines"))
    for number, finding in enumerate(grouped, 1):
        para(doc, f"成果{number}：{finding['vulnerability_family']}", bold=True, style="Heading 2")
        rows = [
            ("序号", str(number).zfill(2)),
            ("成果描述", finding["description"]),
            ("目标系统", finding["system"]),
            ("目标URL", "；".join(finding["urls"]) or "【请补充目标URL】"),
            ("问题类型", finding["vulnerability_family"]),
            ("风险等级", finding["level"] or "待评估"),
        ]
        if finding.get("permission"):
            rows.append(("权限/角色", finding["permission"]))
        rows.extend(optional_scope_rows(finding))
        kv_table(doc, rows)
        _render_reproduction(doc, finding, env_lines)
        _render_result(doc, finding)

    para(doc, "三、存在问题", bold=True, style="Heading 1")
    problems: list[object] = []
    problems.extend(_values(meta.get("problems")))
    for finding in grouped:
        problems.extend(finding.get("problems") or [])
        if not finding.get("problems") and finding.get("interpretation"):
            problems.append(finding["interpretation"])
    for item in _limit(problems, max_items) or ["【请补充本成果对应的核心安全问题】"]:
        para(doc, item)

    para(doc, "四、整改建议", bold=True, style="Heading 1")
    suggestions: list[object] = []
    suggestions.extend(_values(meta.get("suggestions")))
    for finding in grouped:
        suggestions.extend(finding.get("remediations") or [])
    for item in _limit(suggestions, max_items) or ["【请补充与成果对应的整改措施】"]:
        para(doc, item)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


PRIORITY_LEVEL = {"high": "高危", "medium": "中危", "low": "低危"}


def from_ledger(engagement: Path, site: str | None = None) -> tuple[dict, list[dict]]:
    ledger = engagement / "review_ledger.csv"
    with ledger.open(encoding="utf-8-sig", newline="") as f:
        rows = [r for r in csv.DictReader(f) if (r.get("status") or "").strip() == "confirmed"]
    if site:
        rows = [r for r in rows if (r.get("asset") or "").strip().lower() == site.strip().lower()]
    findings = [{
        "finding_id": row.get("finding_id") or row.get("id") or "",
        "title": row.get("summary") or row.get("category") or "待命名成果",
        "desc": row.get("summary") or "", "system": row.get("asset") or "", "url": row.get("endpoint") or "",
        "threat": row.get("category") or "", "level": PRIORITY_LEVEL.get((row.get("priority") or "").strip().lower(), "待定"),
        "permission": row.get("role") or "", "evidence": [row["evidence_ref"]] if row.get("evidence_ref") else [],
    } for row in rows]
    return {"target_name": engagement.name.rsplit("-", 1)[0], "domains": engagement.name, "problems": [], "suggestions": []}, findings


def demo() -> tuple[dict, list]:
    return ({"target_name": "【目标名称】", "domains": "【主域名/小程序标识】", "problems": [], "suggestions": [], "env_lines": []}, [{
        "title": "【漏洞类型】", "desc": "【一句话描述实际影响】", "system": "【目标系统】", "url": "【目标URL或接口】", "threat": "【其他】", "level": "【风险等级】", "steps": [], "commands": [],
    }])


def main() -> int:
    ap = argparse.ArgumentParser(description="攻防成果报告 docx 生成器")
    ap.add_argument("--meta", type=Path)
    ap.add_argument("--findings", type=Path)
    ap.add_argument("--from-ledger", type=Path)
    ap.add_argument("--site")
    ap.add_argument("--demo", action="store_true")
    ap.add_argument("--out", type=Path)
    ap.add_argument("--report-config", type=Path, help="报告阶段专用配置；不读取通用执行配置")
    args = ap.parse_args()
    if args.demo:
        meta, findings = demo(); out = args.out or DEFAULT_TEMPLATE
    elif args.from_ledger:
        meta, findings = from_ledger(args.from_ledger, args.site)
        out = args.out or args.from_ledger / "reports" / f"攻防成果报告_{args.from_ledger.name}{('_' + args.site) if args.site else ''}_{now_date()}.docx"
    elif args.meta and args.findings:
        meta = json.loads(args.meta.read_text(encoding="utf-8")); findings = json.loads(args.findings.read_text(encoding="utf-8"))
        out = args.out or ROOT / f"攻防成果报告_{meta.get('target_name', 'report')}_{now_date()}.docx"
    else:
        ap.error("需要 --demo / --from-ledger / (--meta + --findings) 三者之一")
    config = load_report_config(args.report_config or config_path("reporting"))
    policy = config.get("policy", {}) if isinstance(config.get("policy"), dict) else {}
    template_value = config.get("attack_result_template") or str(DEFAULT_TEMPLATE)
    template_value = template_value.replace("{base}", str(ROOT)).replace("{project_root}", str(ROOT))
    template = Path(template_value)
    if not template.is_absolute(): template = (ROOT / template).resolve()
    path = build_report(meta, findings, out, policy, template)
    print(f"[+] 报告 → {path}")
    print(f"[+] 成果 {len(aggregate_report_findings(findings))} 项；截图由报告人员自行插入")
    return 0


if __name__ == "__main__":
    sys.exit(main())
