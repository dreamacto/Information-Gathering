#!/usr/bin/env python3
"""Read-only DOCX postcheck for report structure and forbidden dynamic content."""
from __future__ import annotations
import argparse
import re
from pathlib import Path
from docx import Document

FORBIDDEN = ("证据截图", "证据：artifacts/", "测试团队：", "报告生成日期：", "执行摘要", "渗透路径", "阶段总结", "综述")
LONG_TITLE = re.compile(r"攻防成果报告.*（.*(?:https?://|\\；|；).*[）)]")

def doc_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    parts.extend(c.text for table in doc.tables for row in table.rows for c in row.cells)
    return "\n".join(parts)

def check(path: Path) -> list[str]:
    text = doc_text(path)
    violations = [f"forbidden text: {item}" for item in FORBIDDEN if item in text]
    if LONG_TITLE.search(text): violations.append("dynamic long title detected")
    if "攻防成果报告" not in text: violations.append("template title missing")
    if not ("复现命令" in text or "验证步骤" in text): violations.append("reproduction section missing")
    return violations

def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    violations = check(args.docx)
    if violations:
        for item in violations: print(f"[!] {item}")
        return 1
    print(f"[+] DOCX postcheck passed: {args.docx}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
