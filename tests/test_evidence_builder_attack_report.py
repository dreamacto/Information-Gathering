import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from evidence_builder import make_attack_result_docx


class AttackResultDocxTests(unittest.TestCase):
    def test_generates_docx_with_team_name_and_screenshot_marker(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            run_dir = root / "run"
            run_dir.mkdir()
            template = root / "template.docx"
            config = root / "config.json"

            doc = Document()
            doc.add_paragraph("template placeholder")
            doc.save(template)

            (run_dir / "confirmed_findings.json").write_text(
                json.dumps(
                    {
                        "findings": [
                            {
                                "system": "测试系统",
                                "target_url": "https://example.test/api/user",
                                "vuln_type": "未授权访问",
                                "description": "接口可读取敏感字段",
                                "exploitability": "登录后可稳定访问敏感字段。",
                                "limitations": "需要授权账号 Cookie。",
                                "remediation": "补充接口鉴权和字段级授权。",
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            config.write_text(
                json.dumps(
                    {
                        "reporting": {
                            "team_name": "观叶识微",
                            "attack_result_template": str(template),
                            "auto_generate_attack_report": True,
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            out = make_attack_result_docx(run_dir, config)

            self.assertIsNotNone(out)
            generated = Document(out)
            text = "\n".join(p.text for p in generated.paragraphs)
            self.assertNotIn("测试团队：", text)
            self.assertNotIn("证据截图", text)
            self.assertIn("未授权访问", text)

    def test_same_asset_and_family_are_one_result_with_optional_data_row(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            run_dir = root / "run"
            run_dir.mkdir()
            config = root / "config.json"
            (run_dir / "confirmed_findings.json").write_text(
                json.dumps({"findings": [
                    {"system": "展示名一", "target_url": "https://example.test/api/a", "vuln_type": "未授权访问", "description": "可读取用户数据", "data_volume": "100条"},
                    {"system": "展示名二", "target_url": "https://example.test/api/b", "vuln_type": "未授权访问", "description": "可读取订单数据", "data_volume": "100条"},
                ]}, ensure_ascii=False), encoding="utf-8"
            )
            config.write_text(json.dumps({"reporting": {"team_name": "测试队", "attack_result_template": "{base}/templates/攻防成果报告_模板.docx", "auto_generate_attack_report": True, "policy": {"screenshot_mode": "manual_insert", "max_problems_or_remediations": 2}}}, ensure_ascii=False), encoding="utf-8")
            out = make_attack_result_docx(run_dir, config)
            generated = Document(out)
            text = "\n".join(p.text for p in generated.paragraphs)
            self.assertEqual(text.count("成果1：未授权访问"), 1)
            table_text = "\n".join(c.text for t in generated.tables for r in t.rows for c in r.cells)
            self.assertIn("https://example.test/api/a；https://example.test/api/b", table_text)
            self.assertIn("涉及数据量", table_text)

    def test_data_row_is_omitted_when_no_data_or_scope_exists(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp); run_dir = root / "run"; run_dir.mkdir(); config = root / "config.json"
            (run_dir / "confirmed_findings.json").write_text(json.dumps({"findings": [{"target_url": "https://example.test/login", "vuln_type": "弱口令登录", "description": "可登录系统"}]}, ensure_ascii=False), encoding="utf-8")
            config.write_text(json.dumps({"reporting": {"team_name": "测试队", "attack_result_template": "{base}/templates/攻防成果报告_模板.docx", "auto_generate_attack_report": True, "policy": {"screenshot_mode": "manual_insert", "max_problems_or_remediations": 2}}}, ensure_ascii=False), encoding="utf-8")
            out = make_attack_result_docx(run_dir, config)
            generated = Document(out)
            table_text = "\n".join(c.text for t in generated.tables for r in t.rows for c in r.cells)
            self.assertNotIn("涉及数据量", table_text)

    def test_different_asset_or_family_remain_separate(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp); run_dir = root / "run"; run_dir.mkdir(); config = root / "config.json"
            rows = [
                {"target_url": "https://one.example/login", "vuln_type": "弱口令登录"},
                {"target_url": "https://one.example/api", "vuln_type": "未授权访问"},
                {"target_url": "https://two.example/login", "vuln_type": "弱口令登录"},
            ]
            (run_dir / "confirmed_findings.json").write_text(json.dumps({"findings": rows}, ensure_ascii=False), encoding="utf-8")
            config.write_text(json.dumps({"reporting": {"team_name": "测试队", "attack_result_template": "{base}/templates/攻防成果报告_模板.docx", "auto_generate_attack_report": True, "policy": {"screenshot_mode": "manual_insert", "max_problems_or_remediations": 2}}}, ensure_ascii=False), encoding="utf-8")
            out = make_attack_result_docx(run_dir, config)
            text = "\n".join(p.text for p in Document(out).paragraphs)
            self.assertEqual(text.count("成果1："), 1)
            self.assertEqual(text.count("成果2："), 1)
            self.assertEqual(text.count("成果3："), 1)


if __name__ == "__main__":
    unittest.main()
