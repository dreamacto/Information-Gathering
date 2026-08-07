from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "项目完整Skill流程与使用手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4CE"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = INK
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = DARK_BLUE
    paragraph.add_run("\n" + body)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    doc.add_paragraph()


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("项目流程手册 | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)
    add_footer(doc)

    doc.add_paragraph("项目完整 Skill 流程与使用手册", style="Title")
    doc.add_paragraph("适用于授权攻防演习、真实环境渗透测试和后续漏洞挖掘的低影响、证据化、分层验证工作流", style="Subtitle")
    add_callout(
        doc,
        "核心原则",
        "默认流程负责发现价值和排除误报；深度验证只对少量已确认目标执行最小化证明。每个关键阶段采用“一主一备”：主工具负责全量，备选工具只做抽样复核、补盲或交叉确认。",
    )

    doc.add_heading("1. 项目定位", level=1)
    doc.add_paragraph(
        "本项目用于授权范围内的安全测试、攻防演习和漏洞挖掘。当前 workflow 层把目标导入、范围校验、低速探测、子域名发现、微信小程序发现、JS/API 深挖、真伪验证、审批闸门和证据报告统一起来，减少工具分散、速率不一致和结果难复核的问题。"
    )
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["gov_exercise_runner.py", "默认主入口：导入目标、检查工具、低速探测、分类、路径检查、JS/API、小程序发现和报告摘要。"],
            ["gov_exercise_workflow.json", "机器可读的阶段、风险、输出、审批闸门。"],
            ["tool_strategy.json", "每阶段“一主一备”的工具策略。"],
            ["wechat_miniapp_discovery.py", "微信小程序/公众号/二维码/搜索 dork 线索生成器，并产出可回流扫描目标。"],
            ["authenticated_session_review.py", "生成登录/注册人工队列，并使用人工提供的 Cookie 做认证态 JS/API 元数据复查。"],
            ["evidence_builder.py", "生成日报草稿、证据索引和平台提交模板。"],
            ["runs/<timestamp>/", "每次执行的目标快照、审计日志、结果、报告和证据目录。"],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("2. 完整主流程", level=1)
    add_numbers(
        doc,
        [
            "范围校验：读取域名或 URL 清单，标准化、去重，只处理授权目标；新发现资产进入待确认/待申请清单。",
            "子域名收集：OneForAll、证书透明度和历史结果优先，新增子域名先探活，确认归属后进入扫描目标。",
            "存活探测：低速 HTTP/HTTPS 元信息获取，记录状态码、标题、Server、Content-Type、跳转、长度和首页 hash。",
            "指纹识别与分类归档：写入 cat_java、cat_net、cat_php、cat_oa、cat_ai、cat_bigscreen、cat_login、cat_api、cat_other。",
            "JS/API 深挖：同站只读抓取首页、robots、sitemap、Swagger/OpenAPI、同站 JS，提取接口、上传、导出、登录、管理类线索。",
            "微信小程序发现：从目标、标题、历史结果和首页/同站 JS 中提取小程序、公众号、二维码和 dork 线索。",
            "小程序目标回流：人工确认主体归属后，把 wechat_subdomain_scan_targets.txt 回流到子域名/存活探测和主流程。",
            "登录注册人工交接：自动生成 manual_auth_queue.csv；由操作员逐一确认可注册入口、登录并提供当前会话 Cookie。",
            "认证态 API 深挖：读取本地 Cookie，分析认证后页面、Webpack/JS 和少量 GET 型 API，只保留结构元数据。",
            "高价值路径发现：检查小型固定路径集，只看存在性和元信息。",
            "真伪验证：对比首页 hash、响应长度、Content-Type、状态码、标题和关键词，排除 SPA 首页、统一错误页和登录跳转。",
            "风险分级与审批闸门：弱口令、SQLMap、命令执行、上传验证、webshell、内网扫描等不进入默认全量流程。",
            "最小化漏洞验证：只对确认有价值的少量目标做低速、可停止的证明性验证。",
            "证据链与报告：生成 evidence_index、daily_report_draft、platform_submission_template、priority_targets 和 run_health。",
        ],
    )

    doc.add_heading("3. 微信小程序流程接入", level=1)
    add_callout(
        doc,
        "放置位置",
        "小程序发现位于 JS/API 深挖之后、高价值路径和漏洞验证之前。它既能用离线历史结果快速生成 dork，也能在授权后低速读取首页和同站 JS，找到 wx_appid、gh_、mp.weixin.qq.com、二维码图片和小程序入口线索。",
        LIGHT_FILL,
    )
    add_table(
        doc,
        ["模式", "触发参数", "是否访问目标", "主要产物"],
        [
            ["离线线索", "--wechat-miniapp", "否", "wechat_unit_keyword_seeds.csv、wechat_search_dorks.txt/csv"],
            ["低速在线", "--wechat-miniapp --wechat-live", "是：首页和同站 JS，受 --delay 控制", "wechat_miniapp_candidates.jsonl、wechat_home_checks.jsonl、wechat_js_checks.jsonl"],
            ["目标回流", "读取 wechat_subdomain_scan_targets.txt", "进入下一轮主流程时才访问", "作为 --targets 输入继续探活和深挖"],
            ["待确认资产", "wechat_pending_extra_assets.txt", "否", "微信平台域名、第三方域名、需人工确认的小程序/公众号线索"],
        ],
        [1.2, 1.6, 1.7, 2.0],
    )
    add_bullets(
        doc,
        [
            "默认不扫描微信平台域名，如 mp.weixin.qq.com、weixin.qq.com、servicewechat.com、wxaurl.cn。",
            "只把已确认属于授权单位或授权源站的 URL 写入 wechat_subdomain_scan_targets.txt。",
            "如果通过二维码或公众号确认出小程序主体、服务域名、业务接口，再把确认后的域名加入下一轮 targets 文件。",
            "小程序线索本身不是漏洞；真正有价值的是小程序后端接口、越权、敏感信息、上传、导出和注册/登录逻辑。",
        ],
    )

    doc.add_heading("4. 登录注册与认证态 API 分支", level=1)
    add_callout(
        doc,
        "适用场景",
        "发现登录页后，由人工完成允许的注册和登录，再把当前会话 Cookie 放入本地会话文件。项目随后复查认证后页面、Webpack/JS、泄露的同主机 API 端口和查询接口，寻找认证后敏感字段结构、文件列表或导出入口。",
        LIGHT_FILL,
    )
    add_table(
        doc,
        ["阶段", "产物/参数", "行为边界"],
        [
            ["登录/注册提示", "manual_auth_queue.csv/json", "自动生成清单，不自动注册、爆破或绕过认证。"],
            ["会话输入", "auth_sessions.local.json", "Cookie 只从本地读取，不写入日志和结果。"],
            ["认证态复查", "--auth-review --auth-cookie-file", "同一授权主机、单并发、GET 型接口、限制 JS 和端点数量。"],
            ["结果", "authenticated_impact_candidates.jsonl", "只记录状态、长度、hash、字段名和风险标签，不保存敏感值。"],
        ],
        [1.35, 2.2, 3.0],
    )
    add_code_block(
        doc,
        '<python.exe> .\\gov_exercise_runner.py --targets "D:\\Desktop\\targets.txt" --resume-run-dir .\\runs\\YYYYMMDD_HHMMSS_gx_gov --auth-review --auth-cookie-file .\\auth_sessions.local.json --auth-max-js 20 --auth-max-endpoints 30 --delay 3',
    )
    add_bullets(
        doc,
        [
            "下载、导出、上传、删除、修改、密码和账号操作只列为人工候选，不自动触发。",
            "如果 JS 暴露同一主机的新 API 端口，可以进入认证态复查；新域名或新子域名仍需先确认授权范围。",
            "发现敏感字段或文件列表后立即停止扩大取证，以最小样本和结构证据证明影响。",
        ],
    )

    doc.add_heading("5. 推荐运行命令", level=1)
    doc.add_paragraph("只创建运行目录和合规材料，不主动探测：")
    add_code_block(doc, '<python.exe> .\\gov_exercise_runner.py --targets "D:\\Desktop\\targets.txt"')
    doc.add_paragraph("低速完整只读流程，包含 JS/API 和小程序离线线索：")
    add_code_block(
        doc,
        '<python.exe> .\\gov_exercise_runner.py --targets "D:\\Desktop\\targets.txt" --probe --fingerprint --high-value-paths --api-discovery --api-confirm --wechat-miniapp --delay 3',
    )
    doc.add_paragraph("低速读取首页和同站 JS 提取小程序线索：")
    add_code_block(
        doc,
        '<python.exe> .\\gov_exercise_runner.py --targets "D:\\Desktop\\targets.txt" --probe --fingerprint --api-discovery --wechat-miniapp --wechat-live --wechat-max-js 3 --delay 3',
    )
    doc.add_paragraph("把小程序阶段确认出的授权域名回流到主流程：")
    add_code_block(
        doc,
        '<python.exe> .\\gov_exercise_runner.py --targets .\\runs\\YYYYMMDD_HHMMSS_gx_gov\\wechat_subdomain_scan_targets.txt --probe --fingerprint --high-value-paths --api-discovery --delay 3',
    )
    doc.add_paragraph("断点续跑：")
    add_code_block(
        doc,
        '<python.exe> .\\gov_exercise_runner.py --targets "D:\\Desktop\\targets.txt" --resume-run-dir .\\runs\\YYYYMMDD_HHMMSS_gx_gov --probe --fingerprint --high-value-paths --api-discovery --api-confirm --wechat-miniapp --delay 3',
    )

    doc.add_heading("6. 每阶段一主一备工具", level=1)
    add_table(
        doc,
        ["步骤", "主工具", "备选/复核工具", "使用方式"],
        [
            ["范围校验", "runner allowlist", "人工复核", "统一口径，避免误扫未授权资产。"],
            ["子域名", "OneForAll", "证书透明度/subfinder/历史结果", "被动来源合并；新增资产先确认归属。"],
            ["存活探测", "runner HTTP probe", "httpx", "httpx 只复核失败和边界结果。"],
            ["指纹识别", "runner 规则", "EHole/TideFinger/P1finger", "复核高价值目标技术栈。"],
            ["JS/API", "api_discovery.py", "katana/PackerFuzzer/API-Explorer", "同站、低速、只读。"],
            ["微信小程序", "wechat_miniapp_discovery.py", "搜索 dork/人工扫码/微信内确认", "脚本生成线索，人工确认主体和授权范围。"],
            ["认证态 API", "authenticated_session_review.py", "浏览器/Burp/Yakit", "人工注册登录并提供 Cookie；脚本只做同主机只读结构复查。"],
            ["高价值路径", "runner 路径集", "浏览器/Burp/Yakit", "只看存在性和元信息。"],
            ["漏洞模板", "nuclei", "afrog", "只对确认候选小范围复核。"],
            ["SQL 注入验证", "sqlmap", "人工请求差异复核", "审批后单点、低 risk/level、加 delay，不导出数据。"],
            ["报告证据", "evidence_builder.py", "人工复核", "统一格式，核对证据和评分。"],
        ],
        [1.25, 1.45, 1.7, 2.1],
    )

    doc.add_heading("7. 速率和边界", level=1)
    add_table(
        doc,
        ["控制项", "默认值", "说明"],
        [
            ["并发", "1", "默认单线程，不并发打真实站点。"],
            ["请求间隔", "2 秒以上", "建议实际演习使用 --delay 3 或更高。"],
            ["同 host 间隔", "2 秒以上", "避免连续请求同一站点。"],
            ["随机抖动", "25%", "避免固定节奏。"],
            ["退避状态码", "429/500/502/503/504", "遇到限流或服务异常自动退避。"],
            ["重复错误停止", "同 host 5 次", "连续异常时停止该 host 当前阶段。"],
        ],
        [1.5, 1.35, 3.65],
    )
    add_callout(
        doc,
        "审批边界",
        "弱口令、爆破、SQLMap、命令执行、上传验证、webshell、内网扫描、数据导出都不是默认全量动作。只有当主流程筛出有价值目标，并明确授权边界、速率、停止条件和证据要求后，才进入最小化验证。",
        CAUTION_FILL,
    )

    doc.add_heading("8. Run 目录关键产物", level=1)
    add_table(
        doc,
        ["文件", "说明"],
        [
            ["targets.csv / targets.json", "本次目标快照。"],
            ["probe_results.jsonl", "存活、标题、跳转、Server、Content-Type、首页 hash。"],
            ["fingerprints.jsonl / cat_*.txt", "指纹和分类归档。"],
            ["api_candidates.jsonl / impact_candidates.jsonl", "JS/API 线索和高价值接口。"],
            ["manual_auth_queue.csv / json", "登录页、认证接口和疑似可注册入口的人工处理队列。"],
            ["authenticated_api_results.jsonl", "认证后页面、JS 和 API 的状态、长度、hash 与 JSON 字段结构。"],
            ["authenticated_impact_candidates.jsonl", "认证后敏感字段结构、文件/导出接口和 source map 候选。"],
            ["wechat_unit_keyword_seeds.csv", "小程序/公众号检索关键词种子。"],
            ["wechat_search_dorks.txt / csv", "快速查找小程序、公众号、扫码入口的 dork。"],
            ["wechat_miniapp_candidates.jsonl", "wx_appid、gh_、微信文章、二维码等候选线索。"],
            ["wechat_subdomain_scan_targets.txt", "可回流到下一轮子域名/存活探测的授权源站目标。"],
            ["wechat_pending_extra_assets.txt", "需确认归属的小程序、微信平台或第三方资产。"],
            ["verified_exposures.jsonl", "真伪验证后较可信的暴露候选。"],
            ["priority_targets.json / reports/priority_review.md", "人工优先复核队列。"],
            ["run_health.json / reports/run_health.md", "覆盖率、误报率和运行质量建议。"],
        ],
        [2.3, 4.2],
    )

    doc.add_heading("9. 明天拿到域名清单后的用法", level=1)
    add_bullets(
        doc,
        [
            "确认目标文件路径、授权范围、是否允许子域名和小程序回流目标。",
            "先跑低速完整只读流程，包含 --wechat-miniapp。",
            "查看 priority_targets.json、impact_candidates.jsonl、wechat_subdomain_scan_targets.txt。",
            "查看 manual_auth_queue.csv；对授权允许注册的目标人工注册登录，再用 Cookie 续跑认证态分支。",
            "把确认属于授权单位的小程序/子域名回流到下一轮 targets。",
            "只对高价值目标做最小化漏洞验证，并保留截图、请求摘要、时间和停止条件。",
        ],
    )

    doc.add_heading("10. 一句话总览", level=1)
    doc.add_paragraph(
        "全量目标先跑低影响发现和真伪验证；登录页进入人工注册/登录队列，Cookie 只用于认证态只读 API 复查；微信小程序阶段扩展关联线索；深度验证只对少量高价值目标执行，并且每一步都有速率、证据、边界和停止条件。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
