#!/usr/bin/env python3
"""Build evidence-oriented report drafts from a controlled exercise run."""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path

from exercise_runtime import now_iso, write_json
from artifact_manifest import create_manifest
from screenshot_queue_builder import build_screenshot_queue
from report_model import aggregate_report_findings, normalize_report_finding, optional_scope_rows

from project_paths import config_path

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_REPORT_CONFIG = config_path("reporting")
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp"}


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def read_jsonl(path: Path) -> list[dict]:
    if not path.exists():
        return []
    rows = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def safe_text(value: object, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def sanitize_filename(value: str) -> str:
    value = re.sub(r'[\\/:*?"<>|]+', "_", value)
    value = re.sub(r"\s+", "_", value).strip("._ ")
    return value or "attack_result"


def list_evidence_files(run_dir: Path) -> list[Path]:
    evidence = run_dir / "evidence"
    if not evidence.exists():
        return []
    return sorted(p for p in evidence.rglob("*") if p.is_file())


def make_evidence_index(run_dir: Path) -> Path:
    files = list_evidence_files(run_dir)
    lines = [
        "# Evidence Index",
        "",
        f"- Generated: {now_iso()}",
        "",
        "| File | Target/Step | Video Time Range | Notes |",
        "| --- | --- | --- | --- |",
    ]
    if files:
        for file in files:
            rel = file.relative_to(run_dir).as_posix()
            lines.append(f"| `{rel}` | TODO | TODO | TODO |")
    else:
        lines.append("| TODO | TODO | TODO | Add screenshots with visible system date/time. |")
    out = run_dir / "reports" / "evidence_index.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def make_daily_report(run_dir: Path) -> Path:
    targets = read_json(run_dir / "targets.json")
    summary = read_json(run_dir / "run_summary.json")
    runtime = read_json(run_dir / "runtime_inventory.json")
    probes = read_jsonl(run_dir / "probe_results.jsonl")
    candidates = read_jsonl(run_dir / "candidate_exposures.jsonl")
    verified = read_jsonl(run_dir / "verified_exposures.jsonl")
    false_positive = read_jsonl(run_dir / "false_positive_exposures.jsonl")
    ok = [p for p in probes if p.get("ok")]
    failed = [p for p in probes if not p.get("ok")]
    missing_tools = summary.get("missing_tools", [])

    lines = [
        "# Attack Daily Draft",
        "",
        f"- Generated: {now_iso()}",
        f"- Run directory: `{run_dir}`",
        f"- Target count: {targets.get('count', 0)}",
        f"- Mode: `{summary.get('mode', 'check')}`",
        "",
        "## Scope And Compliance",
        "",
        "- Approved target list imported and normalized.",
        "- No sensitive data export is recorded in this run directory.",
        "- Screenshot queue is generated under `reports/screenshot_queue.md`; sensitive/authenticated pages still need manual redaction.",
        "",
        "## Runtime Status",
        "",
        f"- Python: `{runtime.get('python') or 'MISSING'}`",
        f"- Java: `{runtime.get('java') or 'MISSING'}`",
        f"- Missing tools: {', '.join(missing_tools) if missing_tools else 'None'}",
        "",
        "## Recon Summary",
        "",
        f"- HTTP probe successes: {len(ok)}",
        f"- HTTP probe failures: {len(failed)}",
        f"- Candidate high-value paths: {len(candidates)}",
        f"- Verified exposures after truth check: {len(verified)}",
        f"- Rejected/weak candidates: {len(false_positive)}",
        "",
    ]

    if ok:
        lines.extend([
            "| URL | Status | Server | Title |",
            "| --- | ---: | --- | --- |",
        ])
        for row in ok[:80]:
            title = str(row.get("title", "")).replace("|", "\\|")
            server = str(row.get("server", "")).replace("|", "\\|")
            lines.append(f"| `{row.get('url')}` | {row.get('status', '')} | {server} | {title} |")
        lines.append("")

    if verified:
        lines.extend([
            "## Verified Exposure Candidates",
            "",
            "| Base URL | Path | Kind | Score | Reasons |",
            "| --- | --- | --- | ---: | --- |",
        ])
        for row in verified[:80]:
            reasons = ", ".join(row.get("verification_reasons", []))
            lines.append(
                f"| `{row.get('base_url')}` | `{row.get('path')}` | {row.get('kind')} | "
                f"{row.get('verification_score', '')} | {reasons} |"
            )
        lines.append("")

    lines.extend([
        "## Findings To Submit",
        "",
        "No finding is auto-claimed by this tool. Add only manually verified, authorized results below.",
        "",
        "| Target | Finding | Permission Level | Evidence | Video Time | Submission Status |",
        "| --- | --- | --- | --- | --- | --- |",
        "| TODO | TODO | TODO | TODO | TODO | Draft |",
        "",
        "## Risk Notes",
        "",
        "- High-risk actions require platform approval before execution.",
        "- Do not modify passwords, delete files, alter data, or store sensitive records.",
        "- Extra assets must be submitted through target application before testing.",
        "- Review `approval_required.md` before any medium/high-risk validation.",
    ])

    out = run_dir / "reports" / "daily_report_draft.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def read_findings_file(path: Path) -> list[dict]:
    if not path.exists():
        return []
    if path.suffix.lower() == ".jsonl":
        return read_jsonl(path)
    data = read_json(path)
    if isinstance(data, list):
        return [row for row in data if isinstance(row, dict)]
    if isinstance(data, dict):
        findings = data.get("findings") or data.get("results") or data.get("items")
        if isinstance(findings, list):
            return [row for row in findings if isinstance(row, dict)]
        return [data]
    return []


def load_confirmed_or_verified_findings(run_dir: Path) -> tuple[list[dict], str]:
    confirmed = []
    for name in ("confirmed_findings.json", "confirmed_findings.jsonl"):
        confirmed.extend(read_findings_file(run_dir / name))
    if confirmed:
        return confirmed, "confirmed_findings"
    verified = read_jsonl(run_dir / "verified_exposures.jsonl")
    if verified:
        return verified, "verified_exposures"
    return [], "none"


def normalize_finding(row: dict, index: int) -> dict:
    base_url = safe_text(row.get("base_url"))
    path = safe_text(row.get("path"))
    url = safe_text(row.get("url") or row.get("target_url") or row.get("endpoint"))
    if not url and base_url:
        url = base_url.rstrip("/") + (path if path.startswith("/") else f"/{path}" if path else "")
    vuln_type = safe_text(
        row.get("vuln_type")
        or row.get("type")
        or row.get("kind")
        or row.get("finding")
        or row.get("risk_type"),
        "待确认漏洞类型",
    )
    reasons = row.get("verification_reasons") or row.get("reasons") or row.get("evidence_reasons") or []
    if isinstance(reasons, list):
        reasons_text = "；".join(safe_text(item) for item in reasons if safe_text(item))
    else:
        reasons_text = safe_text(reasons)
    description = safe_text(
        row.get("description")
        or row.get("title")
        or row.get("summary")
        or f"{vuln_type}：{url}",
        f"成果 {index}",
    )
    exploitability = safe_text(
        row.get("exploitability")
        or row.get("impact")
        or row.get("impact_summary")
        or reasons_text,
        "已发现可验证风险点，需结合截图、响应差异、权限边界和业务影响补充最终可利用性说明。",
    )
    limitations = safe_text(
        row.get("limitations")
        or row.get("limit_conditions")
        or row.get("constraints"),
        "限制条件待补充：需说明账号权限、访问来源、是否依赖登录态、是否只读验证、是否存在 WAF/限速/时间窗口限制。",
    )
    fix = safe_text(
        row.get("fix")
        or row.get("remediation")
        or row.get("suggestion")
        or row.get("recommendation"),
        "修复建议待补充：按漏洞类型补充鉴权、输入校验、最小权限、敏感文件访问控制、日志审计和配置加固。",
    )
    return {
        "index": index,
        "description": description,
        "system": safe_text(row.get("system") or row.get("target_name") or row.get("app_name"), "待补充目标系统"),
        "url": url or "待补充 URL",
        "ip": safe_text(row.get("ip") or row.get("target_ip"), "待补充"),
        "network": safe_text(row.get("network") or row.get("network_area"), "外网"),
        "vuln_type": vuln_type,
        "score": safe_text(row.get("score") or row.get("verification_score") or row.get("expected_score"), "待评估"),
        "process": safe_text(row.get("process") or row.get("attack_process") or row.get("proof"), description),
        "exploitability": exploitability,
        "limitations": limitations,
        "fix": fix,
        "screenshot_desc": safe_text(
            row.get("screenshot_desc") or row.get("screenshot_needed"),
            f"{url or description} 的漏洞证明、时间、登录态/权限边界和关键响应差异",
        ),
    }


def screenshot_files_for_finding(run_dir: Path, finding: dict, all_images: list[Path], total_findings: int) -> list[Path]:
    if not all_images:
        return []
    needle_parts = [
        safe_text(finding.get("url")).lower(),
        safe_text(finding.get("system")).lower(),
        safe_text(finding.get("vuln_type")).lower(),
    ]
    host_tokens = []
    for part in needle_parts:
        host_tokens.extend(re.findall(r"[a-z0-9][a-z0-9.-]+\.[a-z]{2,}", part))
    matches = []
    for image in all_images:
        name = image.name.lower()
        if any(token and token in name for token in host_tokens):
            matches.append(image)
    if matches:
        return matches[:3]
    if total_findings == 1:
        return all_images[:5]
    return []


def add_paragraph(doc, text: str = "", *, style: str | None = None, bold: bool = False, color: str | None = None):
    try:
        paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except KeyError:
        paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def set_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_text(text))
    run.bold = bold


def add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, (key, value) in enumerate(rows):
        set_cell(table.rows[idx].cells[0], key, bold=True)
        set_cell(table.rows[idx].cells[1], value)


def clear_template_body(doc) -> None:
    from docx.oxml.ns import qn

    body = doc._body._element
    for element in list(body):
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def load_report_config(path: Path) -> dict:
    """Load report-only configuration; legacy wrapper is accepted at this boundary."""
    cfg = read_json(path)
    if isinstance(cfg.get("reporting"), dict):
        return cfg["reporting"]
    return cfg


def make_attack_result_docx(run_dir: Path, config_path: Path, *, force: bool = False, skip: bool = False) -> Path | None:
    reporting = load_report_config(config_path)
    policy = reporting.get("policy", {})
    if skip:
        return None
    if not force and not reporting.get("auto_generate_attack_report", True):
        return None

    raw_findings, source = load_confirmed_or_verified_findings(run_dir)
    if not raw_findings:
        return None

    # Keep the automatic run entry point aligned with the report-only renderer.
    from report_docx import build_report
    from report_model import aggregate_report_findings
    template_value = safe_text(reporting.get("attack_result_template"))
    template_value = template_value.replace("{base}", str(BASE_DIR)).replace("{project_root}", str(BASE_DIR))
    template_path = Path(template_value) if template_value else BASE_DIR / "templates" / "攻防成果报告_模板.docx"
    if not template_path.is_absolute():
        template_path = (BASE_DIR / template_path).resolve()
    policy = reporting.get("policy", {}) if isinstance(reporting.get("policy"), dict) else {}
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = run_dir / "reports" / f"攻击成果_{stamp}.docx"
    meta = {"env_lines": [], "problems": [], "suggestions": []}
    build_report(meta, raw_findings, out, policy, template_path)
    return out


def _unique_report_text(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        value = safe_text(value)
        if value and value not in seen:
            seen.add(value)
            result.append(value)
    return result

def make_platform_submission_template(run_dir: Path) -> Path:
    data = {
        "generated_at": now_iso(),
        "report_title": "TODO: target + vulnerability + impact",
        "target_name": "TODO",
        "target_url": "TODO",
        "defense_unit": "TODO",
        "asset_belongs_to_scope": "TODO",
        "result_type": "attack_result",
        "path_nodes": [
            {
                "node_name": "Initial Access / Verification",
                "steps": [
                    {
                        "step_title": "TODO",
                        "operation_summary": "TODO",
                        "evidence_files": [],
                        "video_time_range": "TODO",
                        "system_datetime_visible": False,
                        "sensitive_data_exported": False,
                    }
                ],
            }
        ],
        "approval_notes": {
            "high_risk_approval_id": "",
            "attack_resource_reported": True,
        },
    }
    out = run_dir / "reports" / "platform_submission_template.json"
    write_json(out, data)
    return out


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build evidence report drafts from a run directory")
    parser.add_argument("run_dir", type=Path)
    parser.add_argument("--report-config", type=Path, default=DEFAULT_REPORT_CONFIG)
    parser.add_argument("--attack-report", action="store_true", help="Force attack-result DOCX generation when findings exist")
    parser.add_argument("--no-attack-report", action="store_true", help="Skip attack-result DOCX generation")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    run_dir = args.run_dir.resolve()
    if not run_dir.exists():
        raise SystemExit(f"Run directory does not exist: {run_dir}")
    evidence_index = make_evidence_index(run_dir)
    daily_report = make_daily_report(run_dir)
    submission = make_platform_submission_template(run_dir)
    screenshot_queue = build_screenshot_queue(run_dir)
    attack_report = make_attack_result_docx(
        run_dir,
        args.report_config.resolve(),
        force=args.attack_report,
        skip=args.no_attack_report,
    )
    manifest = create_manifest(run_dir)
    print(json.dumps({
        "evidence_index": str(evidence_index),
        "daily_report": str(daily_report),
        "submission_template": str(submission),
        "screenshot_queue": screenshot_queue,
        "attack_result_docx": str(attack_report) if attack_report else "",
        "artifact_manifest": str(run_dir / "artifact_manifest.json"),
        "artifact_root_sha256": manifest["root_sha256"],
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
